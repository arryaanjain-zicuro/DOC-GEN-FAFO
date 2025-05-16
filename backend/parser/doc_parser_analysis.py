import os, json, re, time
import docx
import openpyxl
from typing import List, Dict, Any
from dotenv import load_dotenv
import google.generativeai as genai
from app.core.gemini_rate_limiter import GeminiRateLimiter

load_dotenv()
genai.configure(api_key=os.getenv("GEMINI_API_KEY"))

model = genai.GenerativeModel("gemini-2.0-flash")
rate_limiter = GeminiRateLimiter(rate_limit_per_minute=20)

# ---------- UTILITY HELPERS ----------

def send_to_gemini(prompt: str, retries: int = 3, key: str = "parsing") -> str:
    for attempt in range(retries):
        try:
            rate_limiter.wait(key)  # Updated to include key
            response = model.generate_content(prompt)
            return response.text
        except Exception as e:
            print(f"[Attempt {attempt + 1}] Gemini error: {e}")
            if attempt == retries - 1:
                raise
            time.sleep(2)


def extract_docx_fields(docx_path: str) -> List[Dict[str, str]]:
    doc = docx.Document(docx_path)
    fields = []

    # Handle "Field: Value" in paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if ":" in text:
            parts = text.split(":", 1)
            label = parts[0].strip()
            value = parts[1].strip()
            if label:
                fields.append({"label": label, "value": value or "No value"})

    # Handle tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if len(cells) >= 2:
                label = cells[0]
                value = cells[1]
                if label:
                    fields.append({"label": label, "value": value or "No value"})

    return fields

def extract_excel_fields(excel_path: str) -> List[Dict[str, str]]:
    wb = openpyxl.load_workbook(excel_path, data_only=True)
    fields = []

    for sheet in wb.sheetnames:
        ws = wb[sheet]
        for row in ws.iter_rows(min_row=1, max_row=30):  # First 30 rows
            label_cell = row[0].value if len(row) > 0 else None
            value_cell = row[1].value if len(row) > 1 else None

            label = str(label_cell).strip() if label_cell else ""
            value = str(value_cell).strip() if value_cell else ""

            if label:
                fields.append({"label": label, "value": value or "No value"})

    return fields
# ---------- MAIN ANALYSIS ----------

def summarize_field_sets(all_docs: List[Dict[str, Any]]) -> Dict[str, Any]:
    """Find common and unique fields."""
    all_field_sets = {
        doc["filename"]: set(field["label"] for field in doc["fields"])
        for doc in all_docs
    }
    all_fields = list(all_field_sets.values())

    common_fields = list(set.intersection(*all_fields)) if len(all_fields) > 1 else []
    unique_fields_per_doc = {
        filename: list(fields - set(common_fields))
        for filename, fields in all_field_sets.items()
    }

    return {
        "common_fields": common_fields,
        "unique_fields_per_doc": unique_fields_per_doc
    }

def build_gemini_prompt(documents: List[Dict[str, Any]]) -> str:
    content = json.dumps(documents, indent=2)
    return f"""
You're analyzing a set of DOCX and XLSX documents.

Each has a list of field names extracted from it. Based on these:

1. Identify fields that are repeated across documents.
2. Suggest which one might be the "alpha" or source template.
3. Provide reasoning based on richness, consistency, or naming patterns.
4. Return JSON:
  {{
    "summary": "...your paragraph...",
    "suggested_alpha": "filename.docx"
  }}

Here is the field data:

{content}
"""

async def analyze_documents(file_paths: List[str]) -> Dict[str, Any]:
    """Main entrypoint for parsing and analyzing a batch of files."""
    documents = []

    for path in file_paths:
        filename = os.path.basename(path)
        ext = os.path.splitext(filename)[1].lower()

        try:
            if ext == ".docx":
                fields = extract_docx_fields(path)
                doc_type = "docx"
            elif ext == ".xlsx":
                fields = extract_excel_fields(path)
                doc_type = "xlsx"
            else:
                continue  # Skip unknown types

            documents.append({
                "filename": filename,
                "type": doc_type,
                "fields": fields
            })
        except Exception as e:
            documents.append({
                "filename": filename,
                "type": "unknown",
                "error": str(e),
                "fields": []
            })

    field_analysis = summarize_field_sets(documents)

    prompt = build_gemini_prompt(documents)
    gemini_raw = send_to_gemini(prompt)

    # Clean response
    cleaned = re.sub(r"^```json\s*|\s*```$", "", gemini_raw.strip())
    try:
        gemini_summary = json.loads(cleaned)
    except Exception as e:
        gemini_summary = {
            "summary": "Gemini failed to parse its own output.",
            "suggested_alpha": None,
            "raw_output": gemini_raw
        }

    return {
        "documents": documents,
        "analysis": field_analysis,
        "gemini_summary": gemini_summary
    }
