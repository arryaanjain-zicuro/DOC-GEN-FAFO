import os
from docx import Document
from openai import OpenAI
from openai.types.chat.chat_completion import ChatCompletion
from openai._exceptions import RateLimitError, OpenAIError 
from dotenv import load_dotenv

load_dotenv()


client = OpenAI(api_key=os.getenv("OPENAI_API_KEY", "your-openai-api-key"))
from typing import List, Dict, Any, Union
import json

# Ensure you set your OpenAI API key via environment variable or hardcode for dev

def extract_tables(doc: Document) -> List[List[Dict[str, str]]]:
    """Extract tables as a list of key-value row dictionaries."""
    tables = []
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            if len(row.cells) >= 2:
                key = row.cells[0].text.strip()
                value = row.cells[1].text.strip()
                if key or value:
                    table_data.append({"key": key, "value": value})
        if table_data:
            tables.append(table_data)
    return tables

def extract_paragraphs(doc: Document) -> List[str]:
    """Extract all non-empty paragraphs."""
    return [para.text.strip() for para in doc.paragraphs if para.text.strip()]

def build_prompt(data: Dict[str, Any]) -> str:
    """Format raw data and build a structured JSON prompt."""
    data_str = json.dumps(data, indent=2)

    prompt = f"""
You are a document analysis AI. Given the following data extracted from a Word document:

{data_str}

Please return a JSON object with:
1. "fields" - A list of extracted fields. Each field must have:
   - "name": The field name (like Issuer, Face Value)
   - "type": Type of data (text, date, currency, number)
   - "source": 'table' or 'paragraph'
   - "value": The actual content/value
2. "relationships" - A list describing logical relationships or hierarchies between fields.
3. "missing_fields" - A list of potentially useful fields that are not found but are typically expected in such documents.

Return **only** the JSON.
"""
    return prompt

import time
import openai

# Modified function with retry logic and exponential backoff
def send_to_gpt(prompt: str, retries: int = 5, delay: int = 5) -> ChatCompletion:
    for attempt in range(retries):
        try:
            response = client.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=[{"role": "user", "content": prompt}]
            )
            return response
        except RateLimitError:
            print(f"[Attempt {attempt+1}] Rate limit exceeded, retrying in {delay} seconds...")
            time.sleep(delay)
            delay *= 2  # Exponential backoff
        except OpenAIError as e:
            print(f"[Attempt {attempt+1}] OpenAI error: {e}")
            raise  # Optional: re-raise or return custom error
    raise Exception("Max retries exceeded. Unable to complete the request.")

def parse_alpha_document(docx_file_path: str) -> Dict[str, Any]:
    """Complete parser pipeline for alpha document."""
    doc = Document(docx_file_path)

    extracted = {
        "tables": extract_tables(doc),
        "paragraphs": extract_paragraphs(doc),
    }

    prompt = build_prompt(extracted)
    gpt_analysis = send_to_gpt(prompt)

    return {
        "raw_data": extracted,
        "gpt_analysis": gpt_analysis
    }

async def test_openai():
    try:
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "You are a helpful assistant."},
                {"role": "user", "content": "Say hello."}
            ]
        )
        return {"status": "success", "response": response.choices[0].message.content}
    except Exception as e:
        return {"status": "error", "error": str(e)}

# For CLI or script usage
# if __name__ == "__main__":
#     import sys
#     from pprint import pprint

#     if len(sys.argv) != 2:
#         print("Usage: python alpha_doc_parser.py path_to_document.docx")
#         sys.exit(1)

#     docx_path = sys.argv[1]
#     result = parse_alpha_document(docx_path)

#     print("\n=== Raw Extracted Data ===")
#     pprint(result["raw_data"])

#     print("\n=== GPT Structured Analysis ===")
#     pprint(result["gpt_analysis"])
