import os, re, json, time
from docx import Document
from dotenv import load_dotenv
from typing import Dict, Any
import google.generativeai as genai
from app.core.gemini_rate_limiter import GeminiRateLimiter

load_dotenv()

# Initialize Gemini API
genai.configure(api_key=os.getenv("GEMINI_API_KEY"))

model = genai.GenerativeModel("gemini-2.0-flash")

# Set rate-limiting parameters (adjust as needed)
CALLS = 20
RATE_LIMIT = 60  # seconds

# Custom rate limiter for Gemini
rate_limiter = GeminiRateLimiter(rate_limit_per_minute=CALLS)

def send_to_gemini(prompt: str, retries: int = 5) -> str:
    """Send request to Gemini with rate limiting."""
    for attempt in range(retries):
        try:
            rate_limiter.wait()  # Throttle before sending request
            response = model.generate_content(prompt)
            return response.text
        except Exception as e:
            print(f"[Attempt {attempt + 1}] Gemini error: {e}")
            if attempt == retries - 1:
                raise
            time.sleep(2)  # fixed delay between retries

def extract_beta_word(docx_file: str) -> Dict[str, Any]:
    """Extract tables and paragraphs from a beta Word document."""
    doc = Document(docx_file)
    tables = []
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            if len(row.cells) >= 2:
                key = row.cells[0].text.strip()
                value = row.cells[1].text.strip()
                if key or value:
                    table_data.append({"key": key, "value": value})
        if table_data:
            tables.append(table_data)

    paragraphs = [para.text.strip() for para in doc.paragraphs if para.text.strip()]

    return {
        "tables": tables,
        "paragraphs": paragraphs,
    }

def build_comparison_prompt(alpha_analysis: Dict[str, Any], beta_data: Dict[str, Any]) -> str:
    """Construct the comparison prompt for Gemini."""
    alpha_fields = json.dumps(alpha_analysis['fields'], indent=2)
    beta_raw = json.dumps(beta_data, indent=2)
    return f"""
You are an AI system that compares an ALPHA document to a BETA document.

Alpha (base template) fields:
{alpha_fields}

Beta (transformed version) raw content:
{beta_raw}

Your tasks:
1. Identify which ALPHA fields are reused in BETA.
2. Explain how each reused field is transformed or formatted (e.g., casing, math operation, string template).
3. Detect any new fields in BETA not derived from ALPHA.
4. Return a JSON with:
    - "field_mappings": List of mappings {{"alpha_field": ..., "beta_occurrence": ..., "transformation": ...}}
    - "unmatched_beta_fields": List of raw beta values not linked to alpha
    - "transformation_notes": Any pattern notes (e.g., formulas, formatting)
"""

def analyze_beta_against_alpha(alpha_analysis: Dict[str, Any], beta_data: Dict[str, Any]) -> Dict[str, Any]:
    """Send Gemini prompt with alpha fields and beta data to analyze patterns."""
    prompt = build_comparison_prompt(alpha_analysis, beta_data)
    gemini_response = send_to_gemini(prompt)
    cleaned_response = re.sub(r"^```json\s*|\s*```$", "", gemini_response.strip())

    # Now safely parse the JSON
    try:
        parsed_result = json.loads(cleaned_response)
    except json.JSONDecodeError as e:
        raise ValueError(f"Invalid JSON in Gemini response: {e}")
    
    return parsed_result

def parse_beta_document(beta_path: str, alpha_analysis: Dict[str, Any]) -> Dict[str, Any]:
    """End-to-end parse of beta document and comparison with alpha."""
    beta_data = extract_beta_word(beta_path)
    mapping_result = analyze_beta_against_alpha(alpha_analysis, beta_data)

    return {
        "beta_raw": beta_data,
        "mapping_result": mapping_result
    }

# Example usage
# if __name__ == "__main__":
#     import sys
#     from backend.app.parser.alpha_doc_parser import parse_alpha_document

#     if len(sys.argv) != 3:
#         print("Usage: python beta_doc_parser.py alpha.docx beta.docx")
#         sys.exit(1)

#     alpha_path = sys.argv[1]
#     beta_path = sys.argv[2]

#     alpha = parse_alpha_document(alpha_path)["gemini_analysis"]
#     result = parse_beta_document(beta_path, alpha)

#     print("\n=== Beta Mappings ===")
#     print(json.dumps(result["mapping_result"], indent=2))
