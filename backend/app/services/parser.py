from docx import Document
import re

def extract_term_sheet_data(docx_file):
    doc = Document(docx_file)
    data = {}

    # Extract from tables
    for table in doc.tables:
        for row in table.rows:
            if len(row.cells) >= 2:
                key = row.cells[0].text.strip()
                value = row.cells[1].text.strip()

                if "Initial Fixing Date" in key:
                    data["Initial Fixing Date"] = value
                elif "Final Fixing Date" in key:
                    data["Final Fixing Date"] = value
                elif "Issuer" in key:
                    data["Issuer"] = value
                elif "Issue Price" in key:
                    data["Issue Price"] = value
                elif "Face Value" in key:
                    data["Face Value"] = value
                elif "Participation Rate" in key:
                    data["Participation Rate"] = value

    # Optional fallback: scan paragraphs for unmatched fields
    def extract_from_paragraphs(pattern, paragraphs):
        for para in paragraphs:
            match = re.search(pattern, para.text, re.IGNORECASE)
            if match:
                return match.group(1).strip()
        return None

    paragraphs = doc.paragraphs

    if "Initial Fixing Date" not in data:
        data["Initial Fixing Date"] = extract_from_paragraphs(
            r"Initial Fixing Date\s*\n*([A-Za-z]+\s+\d{1,2}\s*,\s*\d{4})", paragraphs
        )

    if "Final Fixing Date" not in data:
        data["Final Fixing Date"] = extract_from_paragraphs(
            r"Final Fixing Date\s*\n*([A-Za-z]+\s+\d{1,2}\s*,\s*\d{4})", paragraphs
        )

    return data
