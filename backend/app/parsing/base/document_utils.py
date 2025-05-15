# app/parsing/base/document_utils.py
from docx import Document
from typing import List, Dict, Any

def extract_tables(doc: Document) -> List[List[Dict[str, str]]]:
    tables = []
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            if len(row.cells) >= 2:
                key = row.cells[0].text.strip()
                value = row.cells[1].text.strip()
                if key or value:
                    table_data.append({"key": key, "value": value})
        if table_data:
            tables.append(table_data)
    return tables

def extract_paragraphs(doc: Document) -> List[str]:
    return [para.text.strip() for para in doc.paragraphs if para.text.strip()]

def extract_from_docx(docx_file_path: str) -> Dict[str, Any]:
    doc = Document(docx_file_path)
    return {
        "tables": extract_tables(doc),
        "paragraphs": extract_paragraphs(doc),
    }