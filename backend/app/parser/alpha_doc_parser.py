from docx import Document
import openai
from typing import List, Dict, Any, Union
import json
import os

# Ensure you set your OpenAI API key via environment variable or hardcode for dev
openai.api_key = os.getenv("OPENAI_API_KEY", "your-openai-api-key")

def extract_tables(doc: Document) -> List[List[Dict[str, str]]]:
    """Extract tables as a list of key-value row dictionaries."""
    tables = []
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            if len(row.cells) >= 2:
                key = row.cells[0].text.strip()
                value = row.cells[1].text.strip()
                if key or value:
                    table_data.append({"key": key, "value": value})
        if table_data:
            tables.append(table_data)
    return tables

def extract_paragraphs(doc: Document) -> List[str]:
    """Extract all non-empty paragraphs."""
    return [para.text.strip() for para in doc.paragraphs if para.text.strip()]

def build_prompt(data: Dict[str, Any]) -> str:
    """Format raw data and build a structured JSON prompt."""
    data_str = json.dumps(data, indent=2)

    prompt = f"""
You are a document analysis AI. Given the following data extracted from a Word document:

{data_str}

Please return a JSON object with:
1. "fields" - A list of extracted fields. Each field must have:
   - "name": The field name (like Issuer, Face Value)
   - "type": Type of data (text, date, currency, number)
   - "source": 'table' or 'paragraph'
   - "value": The actual content/value
2. "relationships" - A list describing logical relationships or hierarchies between fields.
3. "missing_fields" - A list of potentially useful fields that are not found but are typically expected in such documents.

Return **only** the JSON.
"""
    return prompt

def send_to_gpt(prompt: str) -> Dict[str, Any]:
    """Send prompt to OpenAI GPT and parse as JSON."""
    response = openai.ChatCompletion.create(
        model="gpt-4",
        messages=[
            {"role": "user", "content": prompt}
        ],
        max_tokens=1200,
        temperature=0.4,
    )

    text_output = response['choices'][0]['message']['content'].strip()
    
    # Attempt to parse JSON output from GPT
    try:
        parsed_output = json.loads(text_output)
    except json.JSONDecodeError:
        raise ValueError("Failed to parse GPT output as JSON:\n" + text_output)
    
    return parsed_output

def parse_alpha_document(docx_file_path: str) -> Dict[str, Any]:
    """Complete parser pipeline for alpha document."""
    doc = Document(docx_file_path)

    extracted = {
        "tables": extract_tables(doc),
        "paragraphs": extract_paragraphs(doc),
    }

    prompt = build_prompt(extracted)
    gpt_analysis = send_to_gpt(prompt)

    return {
        "raw_data": extracted,
        "gpt_analysis": gpt_analysis
    }

# For CLI or script usage
if __name__ == "__main__":
    import sys
    from pprint import pprint

    if len(sys.argv) != 2:
        print("Usage: python alpha_doc_parser.py path_to_document.docx")
        sys.exit(1)

    docx_path = sys.argv[1]
    result = parse_alpha_document(docx_path)

    print("\n=== Raw Extracted Data ===")
    pprint(result["raw_data"])

    print("\n=== GPT Structured Analysis ===")
    pprint(result["gpt_analysis"])
